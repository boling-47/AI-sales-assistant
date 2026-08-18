"""
AI销售分析助手
功能：上传销售Excel -> 销售概况 -> 产品分析 -> 趋势预测 -> 客户分析 -> 地域分析 -> 销售行为分析 -> AI报告
"""
import streamlit as st
from openai import OpenAI
import os
from dotenv import load_dotenv
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from docx import Document
from docx.shared import Inches
from io import BytesIO
import re

load_dotenv()

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

CHART_COLORS = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd',
                '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf']


def safe_write_image(fig, filename):
    try:
        fig.write_image(filename)
    except Exception:
        pass


def style_chart(fig):
    fig.update_layout(
        paper_bgcolor='white',
        plot_bgcolor='white',
        font=dict(color='black', size=14),
        xaxis=dict(tickfont=dict(color='black', size=12), title_font=dict(color='black', size=13)),
        yaxis=dict(tickfont=dict(color='black', size=12), title_font=dict(color='black', size=13)),
        legend=dict(font=dict(color='black', size=12))
    )
    return fig


def extract_province(address):
    if not isinstance(address, str):
        return "未知"
    match = re.match(r'^(.+?[省市])', address)
    return match.group(1) if match else str(address)[:3]


def detect_activity_from_date(date):
    """从日期推断活动类型"""
    month, day = date.month, date.day
    if month == 11 and 1 <= day <= 11:
        return '双11'
    elif month == 12 and 10 <= day <= 12:
        return '双12'
    elif month == 6 and 1 <= day <= 18:
        return '618大促'
    elif month == 1 and 1 <= day <= 7:
        return '新年促销'
    return '无活动'


def detect_holiday_from_date(date):
    """检测节假日"""
    month, day = date.month, date.day
    holidays = [
        (1, (1, 3), '元旦'), (2, (15, 21), '春节'), (4, (4, 6), '清明'),
        (5, (1, 5), '劳动节'), (6, (19, 21), '端午'),
        (9, (25, 27), '中秋'), (10, (1, 7), '国庆'),
    ]
    for h_month, (start, end), name in holidays:
        if month == h_month and start <= day <= end:
            return name
    return None


def get_season(month):
    if 3 <= month <= 5:
        return '春季'
    elif 6 <= month <= 8:
        return '夏季'
    elif 9 <= month <= 11:
        return '秋季'
    else:
        return '冬季'


def get_ai_report(data, forecast_data=None, customer_data=None, geo_data=None, behavior_data=None):
    """调用DeepSeek API获取AI分析报告"""
    if not DEEPSEEK_API_KEY:
        return "⚠️ 未配置DEEPSEEK_API_KEY环境变量，无法使用AI分析功能。"

    prompt = f"""
    你是一名销售分析师。
    根据以下销售数据生成分析报告：
    销售额：{data["总销售额"]}
    销量：{data["总销量"]}
    热销产品：{data["热销产品"]}
    """
    if forecast_data:
        prompt += f"""
    环比变化：{forecast_data.get("环比变化", "无数据")}
    同比变化：{forecast_data.get("同比变化", "无数据")}
    下月预测销量（趋势）：{forecast_data.get("趋势预测", "无数据")}
    下月预测销量（移动平均）：{forecast_data.get("移动平均预测", "无数据")}
    """
    if customer_data:
        prompt += f"""
    --- 客户数据 ---
    客户总数：{customer_data.get("客户总数", "无")}
    活动期间新增客户：{customer_data.get("活动新增客户", "无")}
    客户保持率：{customer_data.get("保持率", "无")}
    客户流失率：{customer_data.get("流失率", "无")}
    高价值客户偏好产品：{customer_data.get("高价值偏好", "无")}
    忠诚客户偏好产品：{customer_data.get("忠诚偏好", "无")}
    价格敏感型客户占比：{customer_data.get("价格敏感占比", "无")}
    品质导向型客户占比：{customer_data.get("品质导向占比", "无")}
    """
    if geo_data:
        prompt += f"""
    --- 地域数据 ---
    覆盖地区数：{geo_data.get("地区数", "无")}
    客户最多地区：{geo_data.get("客户最多地区", "无")}
    价格敏感地区：{geo_data.get("价格敏感地区", "无")}
    品质导向地区：{geo_data.get("品质导向地区", "无")}
    各区域产品偏好：{geo_data.get("区域偏好", "无")}
    """
    if behavior_data:
        prompt += f"""
    --- 销售行为数据 ---
    活动拉动效果：{behavior_data.get("活动拉动", "无")}
    节假日对销量影响：{behavior_data.get("节假日影响", "无")}
    季节性趋势：{behavior_data.get("季节性", "无")}
    价格变动对销量影响：{behavior_data.get("价格影响", "无")}
    """

    prompt += """
    请严格按照下面格式输出：
    ## 📌 销售总结
    总结当前销售情况。
    ## 🔍 核心发现
    列出3条数据发现：
    1.
    2.
    3.
    ## 👥 客户分析
    分析客户保持与流失、高价值与忠诚客户特征、价格敏感型与品质导向型客户差异。
    ## 🗺️ 地域分析
    分析各地区客户规模、稳定性、价格敏感度、产品偏好差异。
    ## 📈 销售行为分析
    分析活动拉动效果、节假日影响、季节性趋势、价格变动对销量的影响。
    ## 💡 优化建议
    列出3条可执行建议：
    1.
    2.
    3.
    要求：
    - 使用中文
    - 简洁专业
    - 像商业报告
    - 如果没有某类数据，跳过对应章节
    """
    try:
        client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"AI分析出错：{e}"


# ============================================================
# 主程序
# ============================================================
st.title("📊 AI销售分析助手")
st.write("上传Excel文件，自动分析销售数据、客户画像、地域分布、销售行为，AI生成报告")

# 侧边栏
with st.sidebar:
    st.header("📋 使用说明")
    st.markdown("""
    **Excel需包含以下列：**
    - 日期
    - 产品
    - 销量
    - 金额
    - 用户ID（可选，用于客户分析）
    - 配送地址（可选，用于地域分析）
    - 活动类型（可选，自动推断已知活动）

    **系统自动生成：**
    - 销售概况与产品排行
    - 同比环比与需求预测
    - 客户保持率/流失率/价值分层
    - 地域偏好与价格敏感度
    - 活动拉动/节假日/季节性分析
    - AI智能分析报告
    """)
    st.markdown("---")
    st.markdown("💡 可下载测试数据体验")

# 下载测试数据
import os as _os
_test_data_path = _os.path.join(_os.path.dirname(__file__), "销售数据_测试.xlsx") if "__file__" in dir() else "销售数据_测试.xlsx"
if _os.path.exists(_test_data_path):
    with open(_test_data_path, "rb") as f:
        st.sidebar.download_button(
            "📥 下载测试数据",
            f.read(),
            file_name="销售数据_测试.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# 上传文件
file = st.file_uploader("请选择Excel文件", type=["xlsx", "xls"])

if file is None:
    st.info("👆 请上传Excel文件，或下载测试数据体验")
    st.stop()

# 读取数据
try:
    df = pd.read_excel(file)
except Exception as e:
    st.error(f"读取文件出错：{e}")
    st.stop()

# 检查必要列
required_cols = ['日期', '产品', '销量', '金额']
missing = [c for c in required_cols if c not in df.columns]
if missing:
    st.error(f"缺少必要列：{missing}")
    st.stop()

# 检查可选列
has_customer_data = '用户ID' in df.columns
has_geo_data = '配送地址' in df.columns
has_activity_col = '活动类型' in df.columns

if has_customer_data:
    st.success("✅ 检测到用户ID列，将生成客户分析")
if has_geo_data:
    st.success("✅ 检测到配送地址列，将生成地域分析")
if has_activity_col:
    st.success("✅ 检测到活动类型列，将生成活动分析")
else:
    st.info("ℹ️ 未检测到活动类型列，系统将从日期自动推断已知活动（双11/618等）")

st.write(f"文件名：{file.name} | 共 {len(df)} 条记录")

# ==============================================
# 数据准备
# ==============================================
df['日期'] = pd.to_datetime(df['日期'])
df['单价'] = (df['金额'] / df['销量']).round(2)
df['月份'] = df['日期'].dt.to_period('M')
df['季节'] = df['日期'].dt.month.apply(get_season)

if not has_activity_col:
    df['活动类型'] = df['日期'].apply(detect_activity_from_date)
df['节假日'] = df['日期'].apply(detect_holiday_from_date)

# ==============================================
# 基础指标
# ==============================================
total_sales = int(df['金额'].sum())
total_volume = int(df['销量'].sum())
product_sales = df.groupby("产品")["销量"].sum().sort_values(ascending=False)
hot_product = product_sales.index[0]

# ==============================================
# 1. 销售概况
# ==============================================
st.subheader("📊 销售分析报告")

st.write("### 💰 销售概况")
col1, col2, col3 = st.columns(3)
with col1:
    st.metric("总销售额", f"¥{total_sales:,}")
with col2:
    st.metric("总销量", f"{total_volume:,} 件")
with col3:
    st.metric("热销产品", hot_product)

if has_customer_data:
    col4, col5 = st.columns(2)
    with col4:
        st.metric("客户总数", f"{df['用户ID'].nunique()} 人")
    with col5:
        avg_customer_spending = int(total_sales / df['用户ID'].nunique())
        st.metric("人均消费", f"¥{avg_customer_spending:,}")

# ==============================================
# 2. 产品分析
# ==============================================
st.write("### 🏆 产品销售分析")
col1, col2 = st.columns(2)

with col1:
    st.write("📦 销量排行")
    ps_df = product_sales.reset_index()
    fig_sales = px.bar(ps_df, x="产品", y="销量", title="产品销量",
                       color="产品", color_discrete_sequence=CHART_COLORS)
    style_chart(fig_sales)
    safe_write_image(fig_sales, "sales_chart.png")
    st.plotly_chart(fig_sales, use_container_width=True)

with col2:
    st.write("💰 金额排行")
    pm_df = df.groupby("产品")["金额"].sum().reset_index().sort_values(by="金额", ascending=False)
    fig_money = px.bar(pm_df, x="产品", y="金额", title="产品销售额",
                       color="产品", color_discrete_sequence=CHART_COLORS)
    style_chart(fig_money)
    safe_write_image(fig_money, "money_chart.png")
    st.plotly_chart(fig_money, use_container_width=True)

st.write("### 🥧 产品销量占比")
fig2 = px.pie(ps_df, names="产品", values="销量", title="产品销售占比",
              color="产品", color_discrete_sequence=CHART_COLORS)
style_chart(fig2)
safe_write_image(fig2, "pie_chart.png")
st.plotly_chart(fig2, use_container_width=True)

# ==============================================
# 3. 销售趋势
# ==============================================
date_col = '日期'
trend = df.groupby(date_col)["销量"].sum().reset_index()

st.write("### 📈 销售趋势")
fig_line = px.line(trend, x=date_col, y="销量", title="销售数量趋势")
style_chart(fig_line)
safe_write_image(fig_line, "line_chart.png")
st.plotly_chart(fig_line, use_container_width=True)

# 按月汇总
monthly_sales = df.groupby('月份')['销量'].sum().sort_index()

# ==============================================
# 4. 同比环比分析
# ==============================================
st.write("### 📊 同比环比分析")
forecast_data = None

if len(monthly_sales) >= 2:
    current_month = monthly_sales.iloc[-1]
    last_month = monthly_sales.iloc[-2]
    mom_change = ((current_month - last_month) / last_month * 100)

    col_m1, col_m2, col_m3 = st.columns(3)
    with col_m1:
        st.metric("本月销量", f"{current_month} 件")
    with col_m2:
        st.metric("上月销量", f"{last_month} 件")
    with col_m3:
        st.metric("环比变化", f"{mom_change:+.1f}%")

    yoy_change = None
    last_year_sales = None
    current_idx = monthly_sales.index[-1]
    last_year_idx = pd.Period(f"{current_idx.year - 1}-{current_idx.month:02d}", freq='M')
    if last_year_idx in monthly_sales.index:
        last_year_sales = monthly_sales.loc[last_year_idx]
        yoy_change = ((current_month - last_year_sales) / last_year_sales * 100)
        col_y1, col_y2, col_y3 = st.columns(3)
        with col_y1:
            st.metric("本月销量", f"{current_month} 件")
        with col_y2:
            st.metric("去年同期销量", f"{last_year_sales} 件")
        with col_y3:
            st.metric("同比变化", f"{yoy_change:+.1f}%")
    else:
        st.info("暂无去年同期数据，无法计算同比")

    monthly_df = monthly_sales.reset_index()
    monthly_df['月份'] = monthly_df['月份'].astype(str)
    fig_monthly = px.bar(monthly_df, x='月份', y='销量', title="月度销量对比",
                         color='月份', color_discrete_sequence=CHART_COLORS)
    style_chart(fig_monthly)
    safe_write_image(fig_monthly, "monthly_chart.png")
    st.plotly_chart(fig_monthly, use_container_width=True)

    forecast_data = {
        "环比变化": f"{mom_change:+.1f}%",
        "同比变化": f"{yoy_change:+.1f}%" if yoy_change is not None else "无去年同期数据",
    }
else:
    st.info("数据不足，至少需要两个月的数据才能做环比分析")

# ==============================================
# 5. 需求预测
# ==============================================
st.write("### 🔮 需求预测")
if len(monthly_sales) >= 3:
    x = np.arange(len(monthly_sales))
    y = monthly_sales.values.astype(float)
    coeffs = np.polyfit(x, y, 1)

    future_labels = []
    trend_preds = []
    last_idx = monthly_sales.index[-1]
    for i in range(1, 4):
        future_x = len(monthly_sales) - 1 + i
        pred = max(int(coeffs[0] * future_x + coeffs[1]), 0)
        trend_preds.append(pred)
        future_labels.append(str(last_idx + i))

    ma_pred = int(monthly_sales.iloc[-3:].mean())

    col_p1, col_p2 = st.columns(2)
    with col_p1:
        st.metric("下月预测（趋势线）", f"{trend_preds[0]} 件")
    with col_p2:
        st.metric("下月预测（移动平均）", f"{ma_pred} 件")

    st.write(f"未来3个月趋势预测：{trend_preds[0]} → {trend_preds[1]} → {trend_preds[2]} 件")

    hist_months = [str(idx) for idx in monthly_sales.index]
    hist_values = list(monthly_sales.values)

    fig_forecast = go.Figure()
    fig_forecast.add_trace(go.Scatter(
        x=hist_months, y=hist_values, mode='lines+markers',
        name='历史销量', line=dict(color='#1f77b4', width=2)
    ))
    fig_forecast.add_trace(go.Scatter(
        x=[hist_months[-1]] + future_labels,
        y=[hist_values[-1]] + trend_preds,
        mode='lines+markers', name='趋势预测',
        line=dict(color='#ff7f0e', width=2, dash='dash'),
        marker=dict(size=10)
    ))
    fig_forecast.update_layout(title="销量趋势与预测")
    style_chart(fig_forecast)
    safe_write_image(fig_forecast, "forecast_chart.png")
    st.plotly_chart(fig_forecast, use_container_width=True)

    st.caption("💡 趋势预测基于线性回归，移动平均基于最近3个月数据。仅供参考，实际备货需结合季节性等因素。")

    if forecast_data:
        forecast_data["趋势预测"] = f"{trend_preds[0]} 件"
        forecast_data["移动平均预测"] = f"{ma_pred} 件"
else:
    st.info("数据不足，至少需要三个月的数据才能做需求预测")

# ==============================================
# 6. 客户分析
# ==============================================
customer_data = None
if has_customer_data:
    st.write("### 👥 客户分析")

    customer_stats = df.groupby('用户ID').agg(
        消费总额=('金额', 'sum'),
        购买次数=('日期', 'count'),
        购买件数=('销量', 'sum'),
        平均单价=('单价', 'mean')
    ).sort_values('消费总额', ascending=False)

    total_customers = len(customer_stats)

    # ---- 6.1 活动期间新增客户与销量增长 ----
    st.write("#### 🎯 活动期间新增客户与销量增长")

    first_purchase = df.groupby('用户ID')['日期'].min()
    activity_new_customers = {}
    activity_uplift = {}

    for activity in df[df['活动类型'] != '无活动']['活动类型'].unique():
        act_mask = df['活动类型'] == activity
        act_start = df[act_mask]['日期'].min()
        act_end = df[act_mask]['日期'].max()

        new_in_activity = sum(
            1 for uid, fdate in first_purchase.items()
            if act_start <= fdate <= act_end
        )
        activity_new_customers[activity] = new_in_activity

        act_daily_avg = df[act_mask]['销量'].sum() / max(1, (act_end - act_start).days + 1)
        before_mask = (df['日期'] < act_start) & (df['日期'] >= act_start - pd.Timedelta(days=30))
        before_daily_avg = df[before_mask]['销量'].sum() / 30 if before_mask.any() else 0
        uplift = ((act_daily_avg - before_daily_avg) / before_daily_avg * 100) if before_daily_avg > 0 else 0
        activity_uplift[activity] = round(uplift, 1)

    if activity_new_customers:
        act_df = pd.DataFrame({
            '活动': list(activity_new_customers.keys()),
            '新增客户数': list(activity_new_customers.values()),
            '销量增长(%)': list(activity_uplift.values())
        })
        col_a1, col_a2 = st.columns(2)
        with col_a1:
            fig_act_new = px.bar(act_df, x='活动', y='新增客户数', title="活动期间新增客户数",
                                 color='活动', color_discrete_sequence=CHART_COLORS)
            style_chart(fig_act_new)
            safe_write_image(fig_act_new, "activity_new_customers.png")
            st.plotly_chart(fig_act_new, use_container_width=True)
        with col_a2:
            fig_act_up = px.bar(act_df, x='活动', y='销量增长(%)', title="活动销量增长(对比前30天日均)",
                                color='活动', color_discrete_sequence=CHART_COLORS)
            style_chart(fig_act_up)
            safe_write_image(fig_act_up, "activity_uplift.png")
            st.plotly_chart(fig_act_up, use_container_width=True)
        st.dataframe(act_df, use_container_width=True, hide_index=True)
    else:
        st.info("未检测到活动期间数据")

    # ---- 6.2 客户保持率与流失率 ----
    st.write("#### 🔄 客户保持率与流失率")

    monthly_customer_sets = df.groupby('月份')['用户ID'].apply(lambda x: set(x.unique()))
    retention_months = []
    retention_rates = []
    churn_rates = []

    for i in range(1, len(monthly_customer_sets)):
        prev_customers = monthly_customer_sets.iloc[i - 1]
        curr_customers = monthly_customer_sets.iloc[i]
        if len(prev_customers) == 0:
            continue
        retained = prev_customers & curr_customers
        churned = prev_customers - curr_customers
        ret_rate = round(len(retained) / len(prev_customers) * 100, 1)
        chn_rate = round(len(churned) / len(prev_customers) * 100, 1)
        retention_months.append(str(monthly_customer_sets.index[i]))
        retention_rates.append(ret_rate)
        churn_rates.append(chn_rate)

    if retention_months:
        col_r1, col_r2 = st.columns(2)
        with col_r1:
            st.metric("平均保持率", f"{np.mean(retention_rates):.1f}%")
        with col_r2:
            st.metric("平均流失率", f"{np.mean(churn_rates):.1f}%")

        retention_df = pd.DataFrame({
            '月份': retention_months,
            '保持率(%)': retention_rates,
            '流失率(%)': churn_rates
        })
        fig_ret = make_subplots(specs=[[{"secondary_y": True}]])
        fig_ret.add_trace(
            go.Scatter(x=retention_months, y=retention_rates, name='保持率(%)',
                       line=dict(color='#27AE60', width=2), mode='lines+markers'),
            secondary_y=False
        )
        fig_ret.add_trace(
            go.Scatter(x=retention_months, y=churn_rates, name='流失率(%)',
                       line=dict(color='#E74C3C', width=2), mode='lines+markers'),
            secondary_y=True
        )
        fig_ret.update_layout(title="客户保持率与流失率趋势")
        fig_ret.update_yaxes(title_text="保持率(%)", secondary_y=False)
        fig_ret.update_yaxes(title_text="流失率(%)", secondary_y=True)
        style_chart(fig_ret)
        safe_write_image(fig_ret, "retention_churn.png")
        st.plotly_chart(fig_ret, use_container_width=True)
    else:
        st.info("数据不足，至少需要两个月才能计算保持率/流失率")

    # ---- 6.3 高价值客户与偏好产品 ----
    st.write("#### 💎 高价值客户与偏好产品")

    top_20_pct = max(int(total_customers * 0.2), 1)
    high_value_customers = customer_stats.head(top_20_pct)
    high_value_revenue = int(high_value_customers['消费总额'].sum())
    high_value_pct = round(high_value_revenue / total_sales * 100, 1)

    hv_products = df[df['用户ID'].isin(high_value_customers.index)].groupby('产品')['销量'].sum().sort_values(ascending=False)

    col_hv1, col_hv2 = st.columns(2)
    with col_hv1:
        st.metric("高价值客户数", f"{top_20_pct} 人")
        st.metric("贡献销售额占比", f"{high_value_pct}%")
    with col_hv2:
        st.metric("高价值客户人均消费", f"¥{int(high_value_customers['消费总额'].mean()):,}")

    fig_hv_prod = px.bar(
        x=hv_products.index, y=hv_products.values,
        title="高价值客户偏好产品", color=hv_products.index,
        color_discrete_sequence=CHART_COLORS
    )
    style_chart(fig_hv_prod)
    fig_hv_prod.update_layout(showlegend=False, xaxis_title="产品", yaxis_title="销量")
    safe_write_image(fig_hv_prod, "high_value_products.png")
    st.plotly_chart(fig_hv_prod, use_container_width=True)

    # ---- 6.4 忠诚客户与偏好产品 ----
    st.write("#### 🤝 忠诚客户与偏好产品")

    freq_sorted = customer_stats.sort_values('购买次数', ascending=False)
    loyal_count = max(int(total_customers * 0.2), 1)
    loyal_customers = freq_sorted.head(loyal_count)

    loyal_products = df[df['用户ID'].isin(loyal_customers.index)].groupby('产品')['销量'].sum().sort_values(ascending=False)

    col_l1, col_l2 = st.columns(2)
    with col_l1:
        st.metric("忠诚客户数", f"{loyal_count} 人")
    with col_l2:
        st.metric("忠诚客户平均购买次数", f"{loyal_customers['购买次数'].mean():.1f} 次")

    fig_loyal_prod = px.bar(
        x=loyal_products.index, y=loyal_products.values,
        title="忠诚客户偏好产品", color=loyal_products.index,
        color_discrete_sequence=CHART_COLORS
    )
    style_chart(fig_loyal_prod)
    fig_loyal_prod.update_layout(showlegend=False, xaxis_title="产品", yaxis_title="销量")
    safe_write_image(fig_loyal_prod, "loyal_products.png")
    st.plotly_chart(fig_loyal_prod, use_container_width=True)

    # ---- 6.5 价格敏感型 vs 产品质量型客户 ----
    st.write("#### 🏷️ 价格敏感型 vs 产品质量型客户")

    price_sorted = customer_stats.sort_values('平均单价')
    low_count = max(int(total_customers * 0.3), 1)
    high_count = max(int(total_customers * 0.3), 1)

    price_sensitive = price_sorted.head(low_count)
    quality_focused = price_sorted.tail(high_count)

    ps_products = df[df['用户ID'].isin(price_sensitive.index)].groupby('产品')['销量'].sum().sort_values(ascending=False)
    qf_products = df[df['用户ID'].isin(quality_focused.index)].groupby('产品')['销量'].sum().sort_values(ascending=False)

    col_ps1, col_ps2 = st.columns(2)
    with col_ps1:
        st.metric("价格敏感型客户", f"{low_count} 人")
        st.caption(f"平均单价 ¥{price_sensitive['平均单价'].mean():.0f}")
    with col_ps2:
        st.metric("品质导向型客户", f"{high_count} 人")
        st.caption(f"平均单价 ¥{quality_focused['平均单价'].mean():.0f}")

    col_pq1, col_pq2 = st.columns(2)
    with col_pq1:
        fig_ps = px.bar(
            x=ps_products.index, y=ps_products.values,
            title="价格敏感型客户偏好产品", color=ps_products.index,
            color_discrete_sequence=CHART_COLORS
        )
        style_chart(fig_ps)
        fig_ps.update_layout(showlegend=False, xaxis_title="产品", yaxis_title="销量")
        safe_write_image(fig_ps, "price_sensitive_products.png")
        st.plotly_chart(fig_ps, use_container_width=True)
    with col_pq2:
        fig_qf = px.bar(
            x=qf_products.index, y=qf_products.values,
            title="品质导向型客户偏好产品", color=qf_products.index,
            color_discrete_sequence=CHART_COLORS
        )
        style_chart(fig_qf)
        fig_qf.update_layout(showlegend=False, xaxis_title="产品", yaxis_title="销量")
        safe_write_image(fig_qf, "quality_focused_products.png")
        st.plotly_chart(fig_qf, use_container_width=True)

    # 客户分析汇总数据
    act_new_str = ", ".join([f"{k}:{v}人" for k, v in activity_new_customers.items()]) if activity_new_customers else "无"
    customer_data = {
        "客户总数": f"{total_customers}人",
        "活动新增客户": act_new_str,
        "保持率": f"{np.mean(retention_rates):.1f}%" if retention_rates else "无",
        "流失率": f"{np.mean(churn_rates):.1f}%" if churn_rates else "无",
        "高价值偏好": ", ".join([f"{p}({int(v)})" for p, v in hv_products.head(3).items()]),
        "忠诚偏好": ", ".join([f"{p}({int(v)})" for p, v in loyal_products.head(3).items()]),
        "价格敏感占比": f"{round(low_count/total_customers*100, 1)}%",
        "品质导向占比": f"{round(high_count/total_customers*100, 1)}%",
    }
else:
    st.write("### 👥 客户分析")
    st.info("当前数据没有「用户ID」列，无法生成客户分析。请在Excel中添加「用户ID」列体验此功能。")

# ==============================================
# 7. 地域分析
# ==============================================
geo_data = None
if has_geo_data:
    st.write("### 🗺️ 地域分析")

    df['省份'] = df['配送地址'].apply(extract_province)

    geo_stats = df.groupby('省份').agg(
        销售额=('金额', 'sum'),
        销量=('销量', 'sum'),
        订单数=('日期', 'count'),
        客户数=('用户ID', 'nunique') if has_customer_data else ('日期', 'count'),
        平均单价=('单价', 'mean')
    ).sort_values('销售额', ascending=False)

    total_regions = len(geo_stats)

    # ---- 7.1 地域产品偏好 ----
    st.write("#### 🔥 地域产品偏好")

    cross = df.pivot_table(values='销量', index='省份', columns='产品', aggfunc='sum', fill_value=0)

    fig_heat = px.imshow(cross, title="地区×产品销量热力图",
                         color_continuous_scale='YlOrRd', aspect='auto')
    fig_heat.update_layout(
        xaxis_title="产品", yaxis_title="省份",
        font=dict(color='black', size=12)
    )
    safe_write_image(fig_heat, "geo_heatmap.png")
    st.plotly_chart(fig_heat, use_container_width=True)

    region_top_product = {}
    for region in cross.index:
        top_prod = cross.loc[region].idxmax()
        top_val = int(cross.loc[region].max())
        region_top_product[region] = f"{top_prod}({top_val}件)"

    st.write("**各地区最偏好产品：**")
    pref_df = pd.DataFrame([
        {'省份': k, '最偏好产品': v.split('(')[0], '销量': int(v.split('(')[1].replace('件)', ''))}
        for k, v in region_top_product.items()
    ]).sort_values('销量', ascending=False)
    st.dataframe(pref_df, use_container_width=True, hide_index=True)

    # ---- 7.2 地域客户规模与稳定性 ----
    st.write("#### 📊 地域客户规模与稳定性")

    total_months = df['月份'].nunique()
    region_active_months = df.groupby('省份')['月份'].nunique()
    region_stability = (region_active_months / total_months * 100).round(1)

    stability_df = pd.DataFrame({
        '省份': geo_stats.index,
        '客户数': geo_stats['客户数'].values if has_customer_data else geo_stats['订单数'].values,
        '活跃月份占比(%)': region_stability.reindex(geo_stats.index).values
    })

    col_s1, col_s2 = st.columns(2)
    with col_s1:
        fig_cust = px.bar(
            stability_df.sort_values('客户数', ascending=False),
            x='省份', y='客户数', title="各地区客户数",
            color='客户数', color_continuous_scale='Blues'
        )
        style_chart(fig_cust)
        fig_cust.update_layout(showlegend=False)
        safe_write_image(fig_cust, "geo_customers.png")
        st.plotly_chart(fig_cust, use_container_width=True)
    with col_s2:
        fig_stab = px.bar(
            stability_df.sort_values('活跃月份占比(%)', ascending=False),
            x='省份', y='活跃月份占比(%)', title="各地区客户稳定性(活跃月份占比)",
            color='活跃月份占比(%)', color_continuous_scale='Greens'
        )
        style_chart(fig_stab)
        fig_stab.update_layout(showlegend=False)
        safe_write_image(fig_stab, "geo_stability.png")
        st.plotly_chart(fig_stab, use_container_width=True)

    top_region = stability_df.sort_values('客户数', ascending=False).iloc[0]['省份']
    most_stable = stability_df.sort_values('活跃月份占比(%)', ascending=False).iloc[0]['省份']

    # ---- 7.3 地域价格敏感度 ----
    st.write("#### 💰 地域价格敏感度")

    geo_price = geo_stats[['平均单价']].sort_values('平均单价')
    median_price = geo_price['平均单价'].median()
    price_sensitive_regions = geo_price[geo_price['平均单价'] <= median_price].index.tolist()
    quality_regions = geo_price[geo_price['平均单价'] > median_price].index.tolist()

    geo_price_df = geo_price.reset_index()
    geo_price_df['分类'] = geo_price_df['平均单价'].apply(
        lambda x: '价格敏感区' if x <= median_price else '品质导向区'
    )

    fig_geo_price = px.bar(
        geo_price_df, x='省份', y='平均单价', title="各地区平均单价对比",
        color='分类', color_discrete_map={'价格敏感区': '#3498DB', '品质导向区': '#E74C3C'}
    )
    style_chart(fig_geo_price)
    safe_write_image(fig_geo_price, "geo_price.png")
    st.plotly_chart(fig_geo_price, use_container_width=True)

    st.info(f"💡 **价格敏感区**（均价≤¥{median_price:.0f}）：{', '.join(price_sensitive_regions)}\n"
            f"**品质导向区**（均价>¥{median_price:.0f}）：{', '.join(quality_regions)}")

    # 地域汇总数据
    top5_pref = ", ".join([f"{r}偏好{region_top_product[r]}" for r in geo_stats.head(5).index])
    geo_data = {
        "地区数": f"{total_regions}个",
        "客户最多地区": f"{top_region}",
        "最稳定地区": f"{most_stable}",
        "价格敏感地区": ", ".join(price_sensitive_regions[:3]),
        "品质导向地区": ", ".join(quality_regions[:3]),
        "区域偏好": top5_pref,
    }
else:
    st.write("### 🗺️ 地域分析")
    st.info("当前数据没有「配送地址」列，无法生成地域分析。请在Excel中添加「配送地址」列体验此功能。")

# ==============================================
# 8. 销售行为分析
# ==============================================
st.write("### 📣 销售行为分析")
behavior_data = None

# ---- 8.1 活动拉动作用 ----
st.write("#### 🎉 平台活动拉动作用")

activity_daily = df.groupby('活动类型').agg(
    总销量=('销量', 'sum'),
    总金额=('金额', 'sum'),
    天数=('日期', 'nunique')
).reset_index()
activity_daily['日均销量'] = (activity_daily['总销量'] / activity_daily['天数']).round(1)

if '无活动' in activity_daily['活动类型'].values:
    normal_daily = activity_daily[activity_daily['活动类型'] == '无活动']['日均销量'].values[0]
    activity_daily['销量提升(%)'] = activity_daily.apply(
        lambda row: round((row['日均销量'] - normal_daily) / normal_daily * 100, 1)
        if row['活动类型'] != '无活动' else 0, axis=1
    )
else:
    activity_daily['销量提升(%)'] = 0

fig_act = px.bar(
    activity_daily, x='活动类型', y='日均销量', title="活动vs非活动日均销量",
    color='活动类型', color_discrete_sequence=CHART_COLORS
)
style_chart(fig_act)
safe_write_image(fig_act, "behavior_activity.png")
st.plotly_chart(fig_act, use_container_width=True)

st.dataframe(activity_daily[['活动类型', '日均销量', '销量提升(%)']], use_container_width=True, hide_index=True)

act_summary = ", ".join([
    f"{r['活动类型']}日均{r['日均销量']}件(提升{r['销量提升(%)']}%)"
    for _, r in activity_daily.iterrows() if r['活动类型'] != '无活动'
])

# ---- 8.2 节假日影响 ----
st.write("#### 🏷️ 节假日对销量影响")

holiday_mask = df['节假日'].notna()
if holiday_mask.any():
    holiday_daily = df[holiday_mask].groupby('节假日').agg(
        总销量=('销量', 'sum'),
        天数=('日期', 'nunique')
    ).reset_index()
    holiday_daily['日均销量'] = (holiday_daily['总销量'] / holiday_daily['天数']).round(1)

    normal_mask = (df['活动类型'] == '无活动') & df['节假日'].isna()
    normal_daily_vol = df[normal_mask]['销量'].sum() / max(1, df[normal_mask]['日期'].nunique())
    holiday_daily['对比日常(%)'] = ((holiday_daily['日均销量'] - normal_daily_vol) / normal_daily_vol * 100).round(1)

    fig_hol = px.bar(
        holiday_daily, x='节假日', y='日均销量', title="节假日vs日常日均销量",
        color='节假日', color_discrete_sequence=CHART_COLORS
    )
    style_chart(fig_hol)
    safe_write_image(fig_hol, "behavior_holiday.png")
    st.plotly_chart(fig_hol, use_container_width=True)
    st.dataframe(holiday_daily, use_container_width=True, hide_index=True)

    hol_summary = ", ".join([
        f"{r['节假日']}日均{r['日均销量']}件(对比日常{r['对比日常(%)']:+.1f}%)"
        for _, r in holiday_daily.iterrows()
    ])
else:
    st.info("未检测到节假日数据")
    hol_summary = "无节假日数据"

# ---- 8.3 季节性影响 ----
st.write("#### 🌦️ 季节对产品及销量影响")

season_product = df.groupby(['季节', '产品'])['销量'].sum().reset_index()
season_order = ['春季', '夏季', '秋季', '冬季']
season_product['季节'] = pd.Categorical(season_product['季节'], categories=season_order, ordered=True)
season_product = season_product.sort_values('季节')

fig_season = px.bar(
    season_product, x='季节', y='销量', color='产品', barmode='group',
    title="各季节产品销量对比", color_discrete_sequence=CHART_COLORS
)
style_chart(fig_season)
safe_write_image(fig_season, "behavior_season.png")
st.plotly_chart(fig_season, use_container_width=True)

season_summary = df.groupby('季节')['销量'].sum().reindex(season_order)
peak_season = season_summary.idxmax()
low_season = season_summary.idxmin()
season_str = f"旺季:{peak_season}({int(season_summary.max())}件), 淡季:{low_season}({int(season_summary.min())}件)"

# ---- 8.4 价格变动影响 ----
st.write("#### 📉 价格变动对销量影响")

monthly_price = df.groupby('月份').agg(
    平均单价=('单价', 'mean'),
    总销量=('销量', 'sum')
).reset_index()
monthly_price['月份'] = monthly_price['月份'].astype(str)

if len(monthly_price) >= 2:
    fig_price = make_subplots(specs=[[{"secondary_y": True}]])
    fig_price.add_trace(
        go.Bar(x=monthly_price['月份'], y=monthly_price['平均单价'],
               name='平均单价', marker_color='#9467BD'),
        secondary_y=False
    )
    fig_price.add_trace(
        go.Scatter(x=monthly_price['月份'], y=monthly_price['总销量'],
                   name='总销量', line=dict(color='#E74C3C', width=2), mode='lines+markers'),
        secondary_y=True
    )
    fig_price.update_layout(title="月度平均单价与销量趋势")
    fig_price.update_yaxes(title_text="平均单价(元)", secondary_y=False)
    fig_price.update_yaxes(title_text="总销量(件)", secondary_y=True)
    style_chart(fig_price)
    safe_write_image(fig_price, "behavior_price.png")
    st.plotly_chart(fig_price, use_container_width=True)

    corr = monthly_price['平均单价'].corr(monthly_price['总销量'])
    st.caption(f"💡 单价与销量的相关系数：{corr:.2f}（{'负相关→降价促销量' if corr < 0 else '正相关→高价高销' if corr > 0 else '无明显相关'}）")
    price_str = f"单价与销量相关系数{corr:.2f}"
else:
    st.info("数据不足，至少需要两个月才能分析价格变动影响")
    price_str = "数据不足"

behavior_data = {
    "活动拉动": act_summary if act_summary else "无活动数据",
    "节假日影响": hol_summary,
    "季节性": season_str,
    "价格影响": price_str,
}

# ==============================================
# 9. AI 分析报告
# ==============================================
st.write("### 🤖 AI分析报告")
with st.spinner("AI正在生成分析报告..."):
    ai_report = get_ai_report(
        {"总销售额": f"¥{total_sales:,}", "总销量": f"{total_volume:,}件", "热销产品": hot_product},
        forecast_data, customer_data, geo_data, behavior_data
    )

with st.container(border=True):
    st.markdown(ai_report)

# ==============================================
# 10. 导出Word报告
# ==============================================
def create_word_report(text):
    doc = Document()
    doc.add_heading("销售分析报告", level=1)
    doc.add_heading("📊 销售数据图表", level=2)

    all_images = [
        "sales_chart.png", "money_chart.png", "pie_chart.png",
        "line_chart.png", "monthly_chart.png", "forecast_chart.png",
        "activity_new_customers.png", "activity_uplift.png", "retention_churn.png",
        "high_value_products.png", "loyal_products.png",
        "price_sensitive_products.png", "quality_focused_products.png",
        "geo_heatmap.png", "geo_customers.png", "geo_stability.png", "geo_price.png",
        "behavior_activity.png", "behavior_holiday.png", "behavior_season.png", "behavior_price.png"
    ]
    for img in all_images:
        try:
            doc.add_picture(img, width=Inches(5))
        except Exception:
            pass

    doc.add_heading("📝 AI分析内容", level=2)
    for line in text.split("\n"):
        if line.startswith("##"):
            doc.add_heading(line.replace("#", "").strip(), level=2)
        elif line.strip():
            line = line.replace("**", "")
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


word_file = create_word_report(ai_report)
st.download_button(
    label="📄 导出Word报告",
    data=word_file,
    file_name="销售分析报告.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
